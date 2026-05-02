"""Utility script for reading and writing .docx files
for the baibaiaigc skill.

This script does NOT run the three-step AIGC reduction itself.
It only helps you move text between .docx files and plain text,
so the skill can work on the text while you keep a .docx workflow.

Requirements:
  pip install python-docx

Typical usage:
  1. Put an input .docx file under the workspace root `origin/` directory.
  2. Run this script to extract its plain text.
  3. Use the baibaiaigc skill on the extracted text.
  4. Optionally, write the final text back into a new .docx file
     under the `finish/` directory.
"""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Iterable

try:
    from docx import Document  # type: ignore[import]
except ImportError as exc:  # pragma: no cover - import guard
    raise SystemExit(
        "Missing dependency python-docx. Install it with: pip install python-docx"
    ) from exc


def read_docx_text(path: Path) -> str:
    """Read a .docx file and return its text as paragraphs joined by blank lines."""
    document = Document(str(path))
    paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs]
    # Keep empty lines only when there is actual content separation.
    non_empty_blocks: list[str] = []
    for paragraph in paragraphs:
        if paragraph:
            non_empty_blocks.append(paragraph)
    return "\n\n".join(non_empty_blocks)


def read_docx_paragraphs(path: Path) -> list[str]:
    """Read a .docx file and return non-empty paragraph texts in order."""
    document = Document(str(path))
    paragraphs: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def write_docx_text(lines: Iterable[str], path: Path) -> None:
    """Write an iterable of text blocks into a .docx file.

    Each element in `lines` becomes one paragraph.
    """
    document = Document()
    for block in lines:
        document.add_paragraph(block)
    document.save(str(path))


def write_docx_paragraphs(paragraphs: Iterable[str], path: Path) -> None:
    """Write one paragraph per entry, preserving paragraph boundaries."""
    write_docx_text(paragraphs, path)


def _split_text_into_blocks(text: str) -> list[str]:
    """Split a large text into paragraph blocks using blank lines as separators."""
    blocks: list[str] = []
    current: list[str] = []
    for raw_line in text.splitlines():
        line = raw_line.rstrip("\n")
        if not line.strip():
            if current:
                blocks.append(" ".join(current).strip())
                current = []
            continue
        current.append(line.strip())
    if current:
        blocks.append(" ".join(current).strip())
    return blocks


def _read_paragraphs_file(path: Path) -> list[str]:
    if path.suffix.lower() == ".json":
        data = json.loads(path.read_text(encoding="utf-8"))
        if not isinstance(data, list) or not all(isinstance(item, str) for item in data):
            raise SystemExit("Paragraph json must be a string array.")
        return data
    text = path.read_text(encoding="utf-8")
    return _split_text_into_blocks(text)


def main(argv: list[str] | None = None) -> None:
    parser = argparse.ArgumentParser(description="Simple .docx <-> text helper")
    subparsers = parser.add_subparsers(dest="command", required=True)

    extract_parser = subparsers.add_parser(
        "extract", help="Extract plain text from a .docx file",
    )
    extract_parser.add_argument("input", type=Path, help="Path to input .docx file")

    extract_to_file_parser = subparsers.add_parser(
        "extract-to-file", help="Extract plain text from a .docx file into a UTF-8 text file",
    )
    extract_to_file_parser.add_argument(
        "input", type=Path, help="Path to input .docx file"
    )
    extract_to_file_parser.add_argument(
        "output", type=Path, help="Path to output .txt file"
    )

    extract_paragraphs_parser = subparsers.add_parser(
        "extract-paragraphs",
        help="Extract non-empty paragraphs from a .docx file into a JSON array",
    )
    extract_paragraphs_parser.add_argument(
        "input", type=Path, help="Path to input .docx file"
    )
    extract_paragraphs_parser.add_argument(
        "output", type=Path, help="Path to output .json file"
    )

    build_parser = subparsers.add_parser(
        "build", help="Build a .docx file from a plain text file",
    )
    build_parser.add_argument("input", type=Path, help="Path to input .txt file")
    build_parser.add_argument("output", type=Path, help="Path to output .docx file")

    build_paragraphs_parser = subparsers.add_parser(
        "build-paragraphs",
        help="Build a .docx file from a paragraph JSON array or block text file",
    )
    build_paragraphs_parser.add_argument("input", type=Path, help="Path to paragraph json/txt file")
    build_paragraphs_parser.add_argument("output", type=Path, help="Path to output .docx file")

    args = parser.parse_args(argv)

    if args.command == "extract":
        text = read_docx_text(args.input)
        print(text)
    elif args.command == "extract-to-file":
        text = read_docx_text(args.input)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(text, encoding="utf-8")
    elif args.command == "extract-paragraphs":
        paragraphs = read_docx_paragraphs(args.input)
        args.output.parent.mkdir(parents=True, exist_ok=True)
        args.output.write_text(
            json.dumps(paragraphs, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
    elif args.command == "build":
        text = args.input.read_text(encoding="utf-8")
        blocks = _split_text_into_blocks(text)
        write_docx_text(blocks, args.output)
    elif args.command == "build-paragraphs":
        paragraphs = _read_paragraphs_file(args.input)
        write_docx_paragraphs(paragraphs, args.output)
    else:  # pragma: no cover - argparse guarantees command
        parser.error("Unknown command")


if __name__ == "__main__":  # pragma: no cover - CLI entry point
    main()
